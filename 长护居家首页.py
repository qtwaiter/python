import pandas as pd
from docx import Document

# 路径可能需要根据您的具体环境进行调整
excel_path = 'D:\Documents\inputall.xlsx'
docx_path = 'D:\Documents\长护险居家长者信息表.docx'

# 读取 Excel 文件，指定需要格式化的列
df = pd.read_excel(excel_path, dtype={'联系电话': str, '亲情护理员联系电话': str})

# 清理数据：将 NaN 转换为空字符串
df.fillna('', inplace=True)

# 加载 Word 文档模板
def replace_placeholders(doc, data):
    """用数据行中的数据替换文档中的占位符"""
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{key}}}'
            if placeholder in paragraph.text:
                # 防止NaN和数值格式问题
                value = str(value).strip('.0') if isinstance(value, float) else str(value)
                paragraph.text = paragraph.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f'{{{key}}}' in cell.text:
                        # 防止NaN和数值格式问题
                        value = str(value).strip('.0') if isinstance(value, float) else str(value)
                        cell.text = cell.text.replace(f'{{{key}}}', value)
    return doc

# 为每行数据生成一个个别文档
def create_individual_documents(df):
    output_files = []
    for index, row in df.iterrows():
        doc = Document(docx_path)
        # 替换占位符
        doc = replace_placeholders(doc, row)
        filename = f'D:\Documents\长护居家对象信息首页\{row["姓名"]}-长护居家.docx'
        doc.save(filename)
        output_files.append(filename)
    return output_files

# 执行函数并存储结果
generated_files = create_individual_documents(df)
print(generated_files)
