import tkinter as tk
from tkinter import filedialog
import logging  # 添加: 导入 logging 模块
import markdown
from docx import Document
from docx.shared import Inches
import os
from tkinter.ttk import Progressbar  # 添加: 导入 Progressbar 模块
from tkinter import messagebox  # 添加: 导入 messagebox 模块


def markdown_to_word(md_file_path, output_docx_path):
    # 读取Markdown文件内容
    with open(md_file_path, 'r', encoding='utf-8') as md_file:
        md_content = md_file.read()

    # 将Markdown内容转换为HTML
    html = markdown.markdown(md_content)

    # 创建一个新的Word文档
    doc = Document()

    # 使用BeautifulSoup解析HTML
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')

    # 遍历HTML中的所有标签
    for element in soup.find_all(True):
        if element.name == 'h1':
            doc.add_heading(element.get_text(strip=True), level=0)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(strip=True), level=1)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(strip=True), level=2)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(strip=True), level=3)
        elif element.name == 'h5':
            doc.add_heading(element.get_text(strip=True), level=4)
        elif element.name == 'h6':
            doc.add_heading(element.get_text(strip=True), level=5)
        elif element.name == 'p':
            text = element.get_text(strip=True)
            if text and (not doc.paragraphs or doc.paragraphs[-1].text != text):
                # 创建新段落
                p = doc.add_paragraph()
                
                # 检查是否有strong标签
                strong_elements = element.find_all('strong')
                if strong_elements:
                    # 如果有strong标签，按原样处理
                    for child in element.children:
                        if child.name == 'strong':
                            p.add_run(child.get_text()).bold = True
                        elif isinstance(child, str):
                            p.add_run(child)
                else:
                    # 如果没有strong标签，检查是否包含冒号
                    if ':' in text:
                        colon_index = text.index(':')
                        before_colon = text[:colon_index]
                        after_colon = text[colon_index:]
                        p.add_run(before_colon).bold = True
                        p.add_run(after_colon)
                    else:
                        p.add_run(text)

        elif element.name == 'img':
            img_path = element.get('src')
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4))
        elif element.name == 'ul':
            # 处理无序列表
            p = doc.add_paragraph(style='List Bullet')
            for li in element.find_all('li', recursive=False):
                p.add_run(li.get_text(strip=True) + ' ')
        elif element.name == 'ol':
            # 处理有序列表
            p = doc.add_paragraph(style='List Number')
            for li in element.find_all('li', recursive=False):
                p.add_run(li.get_text(strip=True) + ' ')
        elif element.name == 'blockquote':
            # 处理引用块
            p = doc.add_paragraph(style='Quote')
            p.add_run(element.get_text(strip=True))
        elif element.name == 'pre':
            # 处理代码块
            p = doc.add_paragraph(style='Code')
            p.add_run(element.get_text(strip=True))
        # 增加对表格的支持
        elif element.name == 'table':
            table = doc.add_table(rows=1, cols=len(element.find_all('th', recursive=False)))
            hdr_cells = table.rows[0].cells
            for i, th in enumerate(element.find_all('th', recursive=False)):
                hdr_cells[i].text = th.get_text(strip=True)
            for tr in element.find_all('tr', recursive=False)[1:]:
                row_cells = table.add_row().cells
                for i, td in enumerate(tr.find_all(['td', 'th'], recursive=False)):
                    row_cells[i].text = td.get_text(strip=True)

        # 增加对水平线的支持
        elif element.name == 'hr':
            doc.add_paragraph(style='Heading 1').add_run('\n')

        # 增加对强调（粗体、斜体）的支持
        elif element.name == 'em':
            if doc.paragraphs:
                run = doc.paragraphs[-1].add_run(element.get_text(strip=True))
                run.italic = True
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.italic = True

    # 保存Word文档
    doc.save(output_docx_path)


if __name__ == "__main__":
    # 创建一个Tkinter根窗口，但不显示它
    root = tk.Tk()
    root.withdraw()

    # 打开文件选择对话框
    md_file_path = filedialog.askopenfilename(title="选择Markdown文件", filetypes=[("Markdown files", "*.md")])
    if not md_file_path:
        print("未选择文件，程序退出。")
        exit()

    # 默认输出文件路径为输入文件路径，但扩展名为.docx
    output_docx_path = os.path.splitext(md_file_path)[0] + '.docx'

    # 增加日志记录
    logging.basicConfig(filename='md_to_word.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    logging.info(f'Starting conversion of {md_file_path} to {output_docx_path}')

    # 增加进度条
    progress = Progressbar(root, orient=tk.HORIZONTAL, length=300, mode='determinate')
    progress.pack(pady=20)
    progress['value'] = 20
    root.update_idletasks()

    try:
        markdown_to_word(md_file_path, output_docx_path)
        progress['value'] = 100
        root.update_idletasks()
        messagebox.showinfo("完成", f"文件已成功转换为 {output_docx_path}")
        logging.info(f'Successfully converted {md_file_path} to {output_docx_path}')
    except Exception as e:
        messagebox.showerror("错误", f"转换过程中发生错误: {str(e)}")
        logging.error(f'Error during conversion: {str(e)}')
