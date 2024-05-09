import docx
import pandas as pd
from datetime import datetime

# Load the input Excel file
df = pd.read_excel('input.xlsx')


# Function to replace placeholders in a docx document
def replace_placeholders_in_docx(doc, data):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for key, value in data.items():
                placeholder = '{{{}}}'.format(key)
                if placeholder in text:
                    if pd.isnull(value):
                        # Replace with an empty string if value is NaN
                        replacement = ''
                    elif isinstance(value, datetime):
                        # Format datetime as date only
                        replacement = value.strftime('%Y-%m-%d')
                    elif isinstance(value, (int, float)):
                        # Format integer or float as string without decimal points
                        replacement = str(int(value)) if isinstance(value, float) and value.is_integer() else str(value)
                    else:
                        # Strip leading/trailing whitespace from string values and convert to string
                        replacement = str(value).strip()
                    text = text.replace(placeholder, replacement)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in data.items():
                            placeholder = '{{{}}}'.format(key)
                            if placeholder in text:
                                if pd.isnull(value):
                                    replacement = ''
                                elif isinstance(value, datetime):
                                    replacement = value.strftime('%Y-%m-%d')
                                elif isinstance(value, (int, float)):
                                    replacement = str(int(value)) if isinstance(value,
                                                                                float) and value.is_integer() else str(
                                        value)
                                else:
                                    replacement = str(value).strip()
                                text = text.replace(placeholder, replacement)
                        run.text = text


# Loop through each row in the Excel file
for index, row in df.iterrows():
    # Load the template DocX file for each iteration
    template_doc = docx.Document('template.docx')

    # Replace placeholders in the template with data from this row
    replace_placeholders_in_docx(template_doc, row.to_dict())

    # Save the new document with the person's name
    filename = f'{row["姓名"]}_长护居家.docx'
    template_doc.save(filename)